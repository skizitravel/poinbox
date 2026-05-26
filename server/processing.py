from __future__ import annotations

import sqlite3
import hashlib
import json
import re
import zipfile
import time
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from pypdf import PdfReader

from server.connectors import IncomingAttachment, IncomingEmail, SampleInboxConnector
from server.config import OCR_PROVIDER
from server.db import log
from server.extraction import classify_purchase_order, extract_purchase_order, normalize_date
from server.master_data import format_structured_address, parse_structured_address, run_master_data_reviews

SUPPORTED_ATTACHMENT_EXTENSIONS = {".pdf", ".txt", ".eml", ".xlsx", ".docx"}
COMMON_PRODUCT_WORDS = {
    "a",
    "an",
    "and",
    "assy",
    "assembly",
    "for",
    "of",
    "part",
    "the",
    "with",
}


def import_samples(conn: sqlite3.Connection, sample_dir: Path, storage_dir: Path) -> dict[str, Any]:
    connector = SampleInboxConnector(sample_dir, storage_dir)
    imported = 0
    skipped = 0
    created_pos = 0
    for incoming in connector.fetch_recent():
        existing = conn.execute(
            "SELECT id FROM emails WHERE provider_message_id = ?", (incoming.provider_message_id,)
        ).fetchone()
        if existing:
            skipped += 1
            continue
        email_id = insert_email(conn, incoming)
        attachment_rows = [insert_attachment(conn, email_id, att) for att in incoming.attachments]
        created_pos += process_email(conn, email_id, attachment_rows)
        imported += 1
    return {"imported": imported, "skipped": skipped, "purchase_orders": created_pos}


def insert_email(conn: sqlite3.Connection, incoming: IncomingEmail) -> int:
    cur = conn.execute(
        """
        INSERT INTO emails (provider, provider_message_id, sender, recipients, subject, received_at, body_text)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (
            incoming.provider,
            incoming.provider_message_id,
            incoming.sender,
            incoming.recipients,
            incoming.subject,
            incoming.received_at,
            incoming.body_text,
        ),
    )
    conn.commit()
    return int(cur.lastrowid)


def insert_attachment(conn: sqlite3.Connection, email_id: int, attachment: IncomingAttachment) -> dict[str, Any]:
    extracted_text = ""
    method = "unsupported"
    page_count = None
    if attachment.content_type == "application/pdf" or attachment.filename.lower().endswith(".pdf"):
        extracted_text, page_count = extract_pdf_text(attachment.local_path)
        method = "pdf_text" if extracted_text.strip() else "ocr_unavailable"
        if not extracted_text.strip():
            ocr_text, ocr_method = extract_pdf_with_vision_or_ocr(attachment.local_path)
            if ocr_text.strip():
                extracted_text = ocr_text
                method = ocr_method
    elif attachment.filename.lower().endswith(".txt"):
        extracted_text = attachment.local_path.read_text(encoding="utf-8")
        method = "text"
    elif attachment.filename.lower().endswith(".xlsx"):
        extracted_text = extract_xlsx_text(attachment.local_path)
        method = "excel_text" if extracted_text.strip() else "excel_text_empty"
    elif attachment.filename.lower().endswith(".xls"):
        method = "unsupported_office"
    elif attachment.filename.lower().endswith(".docx"):
        extracted_text = extract_docx_text(attachment.local_path)
        method = "docx_text" if extracted_text.strip() else "docx_text_empty"
    elif attachment.filename.lower().endswith(".doc"):
        method = "unsupported_office"
    sha256_hash = file_sha256(attachment.local_path)
    cur = conn.execute(
        """
        INSERT INTO attachments (email_id, filename, content_type, local_path, extracted_text, extraction_method, sha256_hash, page_count)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            email_id,
            attachment.filename,
            attachment.content_type,
            str(attachment.local_path),
            extracted_text,
            method,
            sha256_hash,
            page_count,
        ),
    )
    conn.commit()
    if method == "ocr_unavailable":
        log(conn, "warning", "PDF had little or no embedded text; OCR is not configured.", email_id, int(cur.lastrowid))
    if method == "unsupported_office":
        log(conn, "warning", "Office attachment type is not supported by the MVP extractor.", email_id, int(cur.lastrowid), {"filename": attachment.filename})
    return {
        "id": int(cur.lastrowid),
        "filename": attachment.filename,
        "extracted_text": extracted_text,
        "extraction_method": method,
        "sha256_hash": sha256_hash,
    }


def file_sha256(path: Path) -> str | None:
    try:
        digest = hashlib.sha256()
        with path.open("rb") as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                digest.update(chunk)
        return digest.hexdigest()
    except Exception:
        return None


def extract_pdf_text(path: Path) -> tuple[str, int | None]:
    try:
        reader = PdfReader(str(path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text, len(reader.pages)
    except Exception:
        return "", None


def extract_xlsx_text(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except Exception:
        return ""
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
        chunks = [f"Workbook: {path.name}"]
        for sheet in workbook.worksheets:
            if sheet.sheet_state != "visible":
                continue
            chunks.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value).strip() for value in row]
                while values and values[-1] == "":
                    values.pop()
                if any(values):
                    chunks.append("\t".join(values))
        return "\n".join(chunks)
    except Exception:
        return ""


def extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except Exception:
        return extract_docx_text_from_zip(path)
    try:
        document = Document(str(path))
        chunks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    chunks.append("\t".join(cells))
        return "\n".join(chunks)
    except Exception:
        return extract_docx_text_from_zip(path)


def extract_docx_text_from_zip(path: Path) -> str:
    try:
        with zipfile.ZipFile(path) as archive:
            xml = archive.read("word/document.xml")
        root = ElementTree.fromstring(xml)
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        chunks: list[str] = []
        for paragraph in root.findall(".//w:p", namespace):
            text = "".join(node.text or "" for node in paragraph.findall(".//w:t", namespace)).strip()
            if text:
                chunks.append(text)
        return "\n".join(chunks)
    except Exception:
        return ""


def extract_pdf_with_vision_or_ocr(path: Path) -> tuple[str, str]:
    provider = (OCR_PROVIDER or "none").lower()
    if provider in {"", "none", "disabled"}:
        return "", "OCR/vision extraction is not configured for this instance."
    return "", f"OCR provider '{provider}' is configured as a scaffold but is not implemented yet."


def process_email(conn: sqlite3.Connection, email_id: int, attachments: list[dict[str, Any]]) -> int:
    email = conn.execute("SELECT * FROM emails WHERE id = ?", (email_id,)).fetchone()
    email_dict = dict(email)
    attachment_text = "\n\n".join(att.get("extracted_text") or "" for att in attachments)
    filenames = [att["filename"] for att in attachments]
    classification = classify_purchase_order(
        email_dict.get("subject") or "",
        email_dict.get("body_text") or "",
        attachment_text,
        filenames,
    )
    conn.execute(
        """
        UPDATE emails
        SET classification = ?, classification_confidence = ?, classification_explanation = ?, processed_at = ?
        WHERE id = ?
        """,
        (
            classification.label,
            classification.confidence,
            classification.explanation,
            datetime.utcnow().isoformat(),
            email_id,
        ),
    )
    conn.commit()
    if classification.label not in {"possible_po", "purchase_order"}:
        log(conn, "info", "Email was not classified as a purchase order.", email_id)
        return 0

    source_text = attachment_text.strip() or email_dict.get("body_text") or ""
    source_attachment = attachments[0] if attachments else None
    prior_examples = find_similar_extraction_examples(conn, None, source_text)
    started = time.perf_counter()
    extraction = extract_purchase_order(
        source_text,
        email_dict,
        source_attachment["filename"] if source_attachment else None,
        None,
        prior_examples,
    )
    extraction_run_id = log_document_extraction_run(
        conn,
        email_id=email_id,
        attachment_id=source_attachment["id"] if source_attachment else None,
        raw_input_text=source_text,
        extraction=extraction,
        latency_ms=elapsed_ms(started),
        success=True,
    )
    apply_cross_references(conn, extraction)
    duplicate = find_duplicate_purchase_order(conn, extraction)
    if duplicate:
        log(
            conn,
            "warning",
            "Duplicate purchase order skipped.",
            email_id,
            source_attachment["id"] if source_attachment else None,
            {
                "existing_purchase_order_id": duplicate["id"],
                "po_number": extraction.get("po_number"),
                "po_revision": extraction.get("po_revision"),
            },
        )
        return 0
    po_id = insert_purchase_order(conn, email_id, source_attachment["id"] if source_attachment else None, extraction)
    link_extraction_run(conn, extraction_run_id, po_id)
    line_ids: list[int] = []
    for line in extraction.get("lines", []):
        line_ids.append(insert_po_line(conn, po_id, extraction.get("po_number"), line))
    store_extraction_field_evidence(conn, po_id, source_text, source_attachment, extraction, line_ids)
    recalculate_po_total(conn, po_id, extracted_total=extraction.get("total_value"))
    create_duplicate_candidates_for_po(conn, po_id)
    run_master_data_reviews(conn, po_id)
    generate_review_tasks_for_po(conn, po_id)
    log(conn, "info", "Purchase order created for review.", email_id, source_attachment["id"] if source_attachment else None, {"po_id": po_id})
    return 1


def log_document_extraction_run(
    conn: sqlite3.Connection,
    *,
    email_id: int | None = None,
    attachment_id: int | None = None,
    purchase_order_id: int | None = None,
    test_document_id: int | None = None,
    raw_input_text: str = "",
    extraction: dict[str, Any] | None = None,
    latency_ms: int | None = None,
    success: bool = True,
    error_message: str | None = None,
) -> int:
    extraction = extraction or {}
    parsed = {key: value for key, value in extraction.items() if not key.startswith("_")}
    cur = conn.execute(
        """
        INSERT INTO document_extraction_runs (
            email_id, attachment_id, purchase_order_id, test_document_id, extraction_method,
            model_name, prompt_version, raw_input_text, raw_output_json, parsed_output_json,
            success, error_message, latency_ms
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            email_id,
            attachment_id,
            purchase_order_id,
            test_document_id,
            extraction.get("_extraction_method") or "rule_based",
            extraction.get("_model_name"),
            extraction.get("_prompt_version"),
            raw_input_text,
            extraction.get("_raw_output_json"),
            json.dumps(parsed),
            1 if success else 0,
            error_message or extraction.get("_error_message"),
            latency_ms,
        ),
    )
    conn.commit()
    return int(cur.lastrowid)


def link_extraction_run(conn: sqlite3.Connection, run_id: int | None, po_id: int) -> None:
    if not run_id:
        return
    conn.execute("UPDATE document_extraction_runs SET purchase_order_id = ? WHERE id = ?", (po_id, run_id))
    conn.commit()


def elapsed_ms(started: float) -> int:
    return int((time.perf_counter() - started) * 1000)


def insert_purchase_order(conn: sqlite3.Connection, email_id: int, attachment_id: int | None, data: dict[str, Any]) -> int:
    order_type_id = default_order_type_id(conn)
    source_type = source_type_for_email(conn, email_id)
    bill_to_structured = normalized_structured_address(data.get("bill_to_address_structured"), data.get("bill_to_address"))
    ship_to_structured = normalized_structured_address(data.get("ship_to_address_structured"), data.get("ship_to_address"))
    bill_to_address = data.get("bill_to_address") or format_structured_address(bill_to_structured)
    ship_to_address = data.get("ship_to_address") or format_structured_address(ship_to_structured)
    cur = conn.execute(
        """
        INSERT INTO purchase_orders (
            email_id, attachment_id, status, customer_company_name, customer_contact_name, bill_to_address,
            ship_to_address, bill_to_address_structured_json, ship_to_address_structured_json,
            po_number, po_revision, date_received, request_date, total_value, currency, source_type, source_sender,
            source_subject, source_attachment_filename, extraction_confidence, extraction_notes, order_type_id,
            field_confidence_json, quote_number, payment_terms, freight_terms
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            email_id,
            attachment_id,
            "Received",
            data.get("customer_company_name"),
            data.get("customer_contact_name"),
            bill_to_address,
            ship_to_address,
            json.dumps(bill_to_structured),
            json.dumps(ship_to_structured),
            data.get("po_number"),
            data.get("po_revision"),
            system_date_received_for_email(conn, email_id),
            normalize_date(data.get("request_date")),
            data.get("total_value"),
            data.get("currency"),
            source_type,
            data.get("source_sender"),
            data.get("source_subject"),
            data.get("source_attachment_filename"),
            data.get("extraction_confidence"),
            data.get("extraction_notes"),
            order_type_id,
            json.dumps(data.get("field_confidence") or {}),
            data.get("quote_number"),
            data.get("payment_terms"),
            data.get("freight_terms"),
        ),
    )
    conn.commit()
    return int(cur.lastrowid)


def system_date_received_for_email(conn: sqlite3.Connection, email_id: int) -> str | None:
    row = conn.execute("SELECT received_at, created_at FROM emails WHERE id = ?", (email_id,)).fetchone()
    if not row:
        return None
    return normalize_date(row["received_at"]) or normalize_date(row["created_at"])


def source_type_for_email(conn: sqlite3.Connection, email_id: int) -> str:
    row = conn.execute("SELECT provider FROM emails WHERE id = ?", (email_id,)).fetchone()
    provider = (row["provider"] if row else "") or ""
    if provider in {"gmail", "outlook"}:
        return "email"
    if provider == "sample":
        return "sample_import"
    return "manual" if provider == "manual" else "unknown"


def normalized_structured_address(value: Any, fallback_text: str | None = None) -> dict[str, str]:
    keys = ["address_line_1", "address_line_2", "address_line_3", "city", "state", "country", "zip_code"]
    if isinstance(value, dict):
        result = {key: str(value.get(key) or "").strip() for key in keys}
        if any(result.values()):
            return result
    return parse_structured_address(fallback_text)


def find_duplicate_purchase_order(conn: sqlite3.Connection, data: dict[str, Any]) -> sqlite3.Row | None:
    po_number = (data.get("po_number") or "").strip()
    if not po_number:
        return None
    revision = (data.get("po_revision") or "").strip()
    customer = (data.get("customer_company_name") or "").strip()
    rows = conn.execute(
        """
        SELECT id, customer_company_name, po_number, po_revision, total_value
        FROM purchase_orders
        WHERE LOWER(TRIM(po_number)) = LOWER(TRIM(?))
          AND COALESCE(NULLIF(TRIM(po_revision), ''), '') = COALESCE(NULLIF(TRIM(?), ''), '')
        ORDER BY id
        """,
        (po_number, revision),
    ).fetchall()
    if not rows:
        return None
    incoming_total = parse_float(data.get("total_value"))
    if customer:
        customer_key = customer.lower()
        for row in rows:
            existing_customer = (row["customer_company_name"] or "").strip().lower()
            existing_total = parse_float(row["total_value"])
            totals_match = incoming_total is None or existing_total is None or abs(incoming_total - existing_total) <= 0.01
            if (not existing_customer or existing_customer == customer_key) and totals_match:
                return row
        return None
    return rows[0]


def parse_float(value: Any) -> float | None:
    try:
        if value in (None, ""):
            return None
        return float(value)
    except (TypeError, ValueError):
        return None


def create_duplicate_candidates_for_po(conn: sqlite3.Connection, po_id: int) -> None:
    po = conn.execute("SELECT * FROM purchase_orders WHERE id = ?", (po_id,)).fetchone()
    if not po:
        return
    po_number = (po["po_number"] or "").strip()
    revision = (po["po_revision"] or "").strip()
    if po_number:
        rows = conn.execute(
            """
            SELECT *
            FROM purchase_orders
            WHERE id != ? AND LOWER(TRIM(po_number)) = LOWER(TRIM(?))
            """,
            (po_id, po_number),
        ).fetchall()
        for row in rows:
            existing_revision = (row["po_revision"] or "").strip()
            same_revision = (revision or "") == (existing_revision or "")
            current_total = parse_float(po["total_value"])
            other_total = parse_float(row["total_value"])
            total_match = current_total is not None and other_total is not None and abs(current_total - other_total) <= 0.01
            if same_revision and total_match:
                match_type = "same_po_number_revision"
                reason = "Same PO number/revision and matching total found."
                score = 1.0
            elif same_revision:
                match_type = "possible_duplicate"
                reason = "Same PO number/revision but material values differ; review before booking."
                score = 0.88
            elif not revision or not existing_revision:
                match_type = "same_po_number_blank_revision"
                reason = "Same PO number with a blank revision on one record; review duplicate risk."
                score = 0.74
            else:
                match_type = "same_po_number_different_revision"
                reason = "Same PO number with a different revision; confirm this is a legitimate revision."
                score = 0.45
            create_duplicate_candidate(conn, po_id, row["id"], match_type, score, reason)
    attachment = conn.execute("SELECT sha256_hash FROM attachments WHERE id = ?", (po["attachment_id"],)).fetchone() if po["attachment_id"] else None
    if attachment and attachment["sha256_hash"]:
        rows = conn.execute(
            """
            SELECT po.id
            FROM attachments a
            JOIN purchase_orders po ON po.attachment_id = a.id
            WHERE po.id != ? AND a.sha256_hash = ?
            """,
            (po_id, attachment["sha256_hash"]),
        ).fetchall()
        for row in rows:
            create_duplicate_candidate(conn, po_id, row["id"], "same_attachment_hash", 0.98, "Same source attachment hash found on another PO.")


def create_duplicate_candidate(conn: sqlite3.Connection, po_id: int, candidate_po_id: int, match_type: str, score: float, reason: str) -> None:
    conn.execute(
        """
        INSERT INTO po_duplicate_candidates (
            purchase_order_id, candidate_purchase_order_id, match_type, match_score, reason, status
        )
        VALUES (?, ?, ?, ?, ?, 'open')
        ON CONFLICT(purchase_order_id, candidate_purchase_order_id, match_type)
        DO UPDATE SET match_score = excluded.match_score, reason = excluded.reason
        """,
        (po_id, candidate_po_id, match_type, score, reason),
    )


def find_similar_extraction_examples(
    conn: sqlite3.Connection,
    customer_company_name: str | None,
    source_text: str | None,
    limit: int = 3,
) -> list[dict[str, Any]]:
    customer = (customer_company_name or "").strip().lower()
    rows = conn.execute(
        """
        SELECT po.id, po.customer_company_name, po.po_number, po.source_attachment_filename,
               po.extraction_notes, po.updated_at
        FROM purchase_orders po
        WHERE po.extraction_feedback_count > 0 OR po.extraction_reviewed_at IS NOT NULL
        ORDER BY
            CASE WHEN LOWER(TRIM(COALESCE(po.customer_company_name, ''))) = ? THEN 0 ELSE 1 END,
            po.updated_at DESC
        LIMIT ?
        """,
        (customer, limit),
    ).fetchall()
    examples = []
    for row in rows:
        feedback = conn.execute(
            """
            SELECT entity_type, field_name, extracted_value, corrected_value, confidence
            FROM extraction_feedback
            WHERE purchase_order_id = ?
            ORDER BY created_at DESC LIMIT 20
            """,
            (row["id"],),
        ).fetchall()
        header = conn.execute(
            """
            SELECT customer_company_name, customer_contact_name, bill_to_address, ship_to_address,
                   po_number, po_revision, quote_number, payment_terms, freight_terms, total_value, currency
            FROM purchase_orders WHERE id = ?
            """,
            (row["id"],),
        ).fetchone()
        lines = conn.execute(
            """
            SELECT line_number, customer_part_number, customer_part_revision, internal_part_number,
                   internal_part_revision, description, quantity, unit_of_measure, unit_price,
                   line_total, requested_date
            FROM purchase_order_lines WHERE purchase_order_id = ? ORDER BY id LIMIT 5
            """,
            (row["id"],),
        ).fetchall()
        examples.append(
            {
                "customer_company_name": row["customer_company_name"],
                "source_attachment_filename": row["source_attachment_filename"],
                "corrected_header_fields": dict(header) if header else {},
                "corrected_line_examples": [dict(line) for line in lines],
                "feedback_rows": [dict(item) for item in feedback],
                "extraction_notes": row["extraction_notes"],
            }
        )
    return examples


def insert_po_line(conn: sqlite3.Connection, purchase_order_id: int, po_number: str | None, line: dict[str, Any]) -> int:
    line_total = normalized_line_total(line.get("quantity"), line.get("unit_price"), line.get("line_total"))
    product_match = match_product_for_line(conn, purchase_order_id, line)
    cur = conn.execute(
        """
        INSERT INTO purchase_order_lines (
            purchase_order_id, po_number, line_number, customer_part_number, customer_part_revision,
            internal_part_number, internal_part_revision, description,
            quantity, unit_of_measure, unit_price, line_total, requested_date, extraction_confidence, extraction_notes,
            field_confidence_json, product_match_status, matched_product_id, product_match_score, product_match_reason,
            canonical_product_id, customer_product_alias_id
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            purchase_order_id,
            line.get("po_number") or po_number,
            line.get("line_number"),
            line.get("customer_part_number"),
            line.get("customer_part_revision"),
            line.get("internal_part_number"),
            line.get("internal_part_revision"),
            line.get("description"),
            line.get("quantity"),
            line.get("unit_of_measure"),
            line.get("unit_price"),
            line_total,
            line.get("requested_date"),
            line.get("extraction_confidence"),
            line.get("extraction_notes"),
            json.dumps(line.get("field_confidence") or {}),
            product_match["status"],
            product_match["product_id"],
            product_match["score"],
            product_match["reason"],
            product_match.get("canonical_product_id"),
            product_match.get("customer_product_alias_id"),
        ),
    )
    conn.commit()
    return int(cur.lastrowid)


def default_order_type_id(conn: sqlite3.Connection) -> int | None:
    row = conn.execute("SELECT id FROM order_types WHERE name = 'Standard' AND is_active = 1").fetchone()
    return row["id"] if row else None


def normalized_line_total(quantity: Any, unit_price: Any, line_total: Any) -> float | None:
    if quantity not in (None, "") and unit_price not in (None, ""):
        return round(float(quantity) * float(unit_price), 2)
    return None if line_total in (None, "") else float(line_total)


def recalculate_po_total(conn: sqlite3.Connection, po_id: int, extracted_total: float | None = None) -> float:
    row = conn.execute(
        "SELECT COALESCE(SUM(COALESCE(line_total, quantity * unit_price)), 0) AS total FROM purchase_order_lines WHERE purchase_order_id = ?",
        (po_id,),
    ).fetchone()
    total = round(float(row["total"] or 0), 2)
    po = conn.execute("SELECT extraction_notes FROM purchase_orders WHERE id = ?", (po_id,)).fetchone()
    notes = po["extraction_notes"] or ""
    if extracted_total not in (None, "") and abs(float(extracted_total) - total) > 0.01:
        warning = f" Header total {extracted_total} differed from calculated line total {total}; review."
        if warning not in notes:
            notes = (notes + warning).strip()
    conn.execute(
        "UPDATE purchase_orders SET total_value = ?, extraction_notes = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
        (total, notes, po_id),
    )
    conn.commit()
    return total


def store_extraction_field_evidence(
    conn: sqlite3.Connection,
    po_id: int,
    source_text: str,
    source_attachment: dict[str, Any] | None,
    extraction: dict[str, Any],
    line_ids: list[int],
) -> None:
    try:
        conn.execute("DELETE FROM extraction_field_evidence WHERE purchase_order_id = ?", (po_id,))
        filename = (source_attachment or {}).get("filename")
        attachment_method = (source_attachment or {}).get("extraction_method") or ""
        header_conf = extraction.get("field_confidence") or {}
        for field in [
            "po_number",
            "po_revision",
            "customer_company_name",
            "customer_contact_name",
            "date_received",
            "quote_number",
            "payment_terms",
            "freight_terms",
            "bill_to_address",
            "ship_to_address",
            "total_value",
        ]:
            value = value_from_evidence_object(extraction.get(field))
            if value in (None, ""):
                continue
            insert_field_evidence(
                conn,
                po_id,
                None,
                field,
                value,
                snippet_from_value(source_text, value),
                filename,
                header_conf.get(field),
                evidence_location_from_source(source_text, value, attachment_method),
            )
        for index, line in enumerate(extraction.get("lines") or []):
            line_id = line_ids[index] if index < len(line_ids) else None
            line_conf = line.get("field_confidence") or {}
            for field in [
                "line_number",
                "customer_part_number",
                "customer_part_revision",
                "internal_part_number",
                "internal_part_revision",
                "description",
                "quantity",
                "unit_of_measure",
                "unit_price",
                "line_total",
                "requested_date",
            ]:
                value = value_from_evidence_object(line.get(field))
                if value in (None, ""):
                    continue
                insert_field_evidence(
                    conn,
                    po_id,
                    line_id,
                    field,
                    value,
                    snippet_from_value(source_text, value),
                    filename,
                    line_conf.get(field),
                    evidence_location_from_source(source_text, value, attachment_method),
                )
        conn.commit()
    except Exception as exc:
        log(conn, "warning", "Extraction field evidence capture failed.", metadata={"po_id": po_id, "error": str(exc)})


def value_from_evidence_object(value: Any) -> Any:
    if isinstance(value, dict):
        return value.get("value")
    return value


def snippet_from_value(source_text: str, value: Any) -> str:
    text = source_text or ""
    needle = str(value or "").strip()
    if not text:
        return ""
    if needle:
        idx = text.lower().find(needle.lower())
        if idx >= 0:
            start = max(0, idx - 160)
            end = min(len(text), idx + len(needle) + 160)
            return text[start:end].strip()
    return text[:360].strip()


def evidence_location_from_source(source_text: str, value: Any, method: str) -> dict[str, Any]:
    needle = str(value or "").strip().lower()
    lines = (source_text or "").splitlines()
    row_number = None
    sheet_name = None
    for index, line in enumerate(lines, start=1):
        stripped = line.strip()
        if stripped.lower().startswith("sheet:"):
            sheet_name = stripped.split(":", 1)[-1].strip()
        if needle and needle in stripped.lower():
            row_number = index
            break
    location: dict[str, Any] = {}
    if method == "excel_text":
        location["sheet_name"] = sheet_name
        location["row_number"] = row_number
    elif method == "docx_text":
        location["paragraph_index"] = row_number
    return location


def insert_field_evidence(
    conn: sqlite3.Connection,
    po_id: int,
    line_id: int | None,
    field_name: str,
    value: Any,
    snippet: str,
    filename: str | None,
    confidence: Any,
    location: dict[str, Any],
) -> None:
    try:
        score = None if confidence in (None, "") else float(confidence)
    except (TypeError, ValueError):
        score = None
    conn.execute(
        """
        INSERT INTO extraction_field_evidence (
            purchase_order_id, purchase_order_line_id, field_name, extracted_value,
            source_snippet, source_attachment_filename, sheet_name, row_number,
            paragraph_index, table_index, email_section, confidence
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            po_id,
            line_id,
            field_name,
            "" if value is None else str(value),
            snippet,
            filename,
            location.get("sheet_name"),
            location.get("row_number"),
            location.get("paragraph_index"),
            location.get("table_index"),
            location.get("email_section"),
            score,
        ),
    )


def apply_cross_references(conn: sqlite3.Connection, extraction: dict[str, Any]) -> None:
    customer = (extraction.get("customer_company_name") or "").strip().lower()
    if not customer:
        return
    for line in extraction.get("lines", []):
        if line.get("internal_part_number"):
            continue
        customer_part = (line.get("customer_part_number") or "").strip().lower()
        if not customer_part:
            continue
        customer_revision = (line.get("customer_part_revision") or "").strip().lower()
        alias = find_canonical_customer_product_alias(conn, customer, customer_part, customer_revision)
        if alias and alias["product_number"]:
            line["internal_part_number"] = alias["product_number"]
            if alias["product_revision"] and not line.get("internal_part_revision"):
                line["internal_part_revision"] = alias["product_revision"]
            notes = line.get("extraction_notes") or ""
            line["extraction_notes"] = (notes + " Matched from canonical customer product alias.").strip()
            confidence = line.setdefault("field_confidence", {})
            confidence["internal_part_number"] = 0.95
            confidence["internal_part_revision"] = 0.95
            continue
        row = conn.execute(
            """
            SELECT internal_part_number, internal_part_revision
            FROM customer_part_xrefs
            WHERE LOWER(TRIM(customer_name)) = ? AND LOWER(TRIM(customer_part_number)) = ?
              AND LOWER(TRIM(COALESCE(customer_part_revision, ''))) = ?
            """,
            (customer, customer_part, customer_revision),
        ).fetchone()
        if not row:
            row = conn.execute(
                """
                SELECT internal_part_number, internal_part_revision
                FROM customer_part_xrefs
                WHERE LOWER(TRIM(customer_name)) = ? AND LOWER(TRIM(customer_part_number)) = ?
                  AND COALESCE(NULLIF(TRIM(customer_part_revision), ''), '') = ''
                """,
                (customer, customer_part),
            ).fetchone()
        if not row:
            continue
        line["internal_part_number"] = row["internal_part_number"]
        if row["internal_part_revision"] and not line.get("internal_part_revision"):
            line["internal_part_revision"] = row["internal_part_revision"]
        notes = line.get("extraction_notes") or ""
        line["extraction_notes"] = (notes + " Matched from customer product cross reference.").strip()
        confidence = line.setdefault("field_confidence", {})
        confidence["internal_part_number"] = 0.95
        confidence["internal_part_revision"] = 0.95


def find_canonical_customer_product_alias(
    conn: sqlite3.Connection, customer_name: str, customer_part_number: str, customer_revision: str | None
) -> sqlite3.Row | None:
    if not customer_name or not customer_part_number:
        return None
    params = (customer_name, customer_part_number, customer_revision or "")
    row = conn.execute(
        """
        SELECT cpa.*, p.product_number
        FROM customer_product_alias cpa
        JOIN trading_partner_account tpa ON tpa.id = cpa.trading_partner_account_id
        LEFT JOIN product p ON p.id = cpa.product_id
        WHERE LOWER(TRIM(tpa.account_name)) = LOWER(TRIM(?))
          AND cpa.normalized_customer_product_number = LOWER(TRIM(?))
          AND LOWER(TRIM(COALESCE(cpa.customer_product_revision, ''))) = LOWER(TRIM(?))
          AND cpa.active_flag = 1
        LIMIT 1
        """,
        params,
    ).fetchone()
    if row:
        return row
    return conn.execute(
        """
        SELECT cpa.*, p.product_number
        FROM customer_product_alias cpa
        JOIN trading_partner_account tpa ON tpa.id = cpa.trading_partner_account_id
        LEFT JOIN product p ON p.id = cpa.product_id
        WHERE LOWER(TRIM(tpa.account_name)) = LOWER(TRIM(?))
          AND cpa.normalized_customer_product_number = LOWER(TRIM(?))
          AND COALESCE(NULLIF(TRIM(cpa.customer_product_revision), ''), '') = ''
          AND cpa.active_flag = 1
        LIMIT 1
        """,
        (customer_name, customer_part_number),
    ).fetchone()


def canonical_product_id_for_legacy_product(conn: sqlite3.Connection, product_id: int | None) -> int | None:
    if not product_id:
        return None
    row = conn.execute(
        "SELECT id FROM product WHERE legacy_source_table = 'products' AND legacy_source_id = ? LIMIT 1",
        (str(product_id),),
    ).fetchone()
    return int(row["id"]) if row else None


def canonical_alias_for_po_line(conn: sqlite3.Connection, po_id: int, line: dict[str, Any]) -> sqlite3.Row | None:
    po = conn.execute("SELECT customer_company_name FROM purchase_orders WHERE id = ?", (po_id,)).fetchone()
    if not po:
        return None
    return find_canonical_customer_product_alias(
        conn,
        po["customer_company_name"] or "",
        (line.get("customer_part_number") or "").strip(),
        (line.get("customer_part_revision") or "").strip(),
    )


def match_product_for_line(conn: sqlite3.Connection, po_id: int, line: dict[str, Any]) -> dict[str, Any]:
    alias = canonical_alias_for_po_line(conn, po_id, line)
    if alias and alias["product_number"] and not line.get("internal_part_number"):
        line["internal_part_number"] = alias["product_number"]
        if alias["product_revision"] and not line.get("internal_part_revision"):
            line["internal_part_revision"] = alias["product_revision"]
    internal_part = (line.get("internal_part_number") or "").strip()
    internal_revision = (line.get("internal_part_revision") or "").strip()
    if internal_part:
        row = conn.execute(
            """
            SELECT * FROM products
            WHERE is_active = 1 AND LOWER(TRIM(internal_part_number)) = LOWER(TRIM(?))
            ORDER BY CASE WHEN LOWER(TRIM(COALESCE(internal_part_revision, ''))) = LOWER(TRIM(?)) THEN 0 ELSE 1 END
            LIMIT 1
            """,
            (internal_part, internal_revision),
        ).fetchone()
        if row:
            return {
                "status": "matched_exact",
                "product_id": row["id"],
                "canonical_product_id": canonical_product_id_for_legacy_product(conn, row["id"]) or (alias["product_id"] if alias else None),
                "customer_product_alias_id": alias["id"] if alias else None,
                "score": 1.0,
                "reason": "Matched active product by internal part number.",
            }
    description = line.get("description") or ""
    best_row = None
    best_score = 0.0
    source_tokens = product_tokens(description)
    if source_tokens:
        for row in conn.execute("SELECT * FROM products WHERE is_active = 1 AND COALESCE(description, '') != ''").fetchall():
            score = token_similarity(source_tokens, product_tokens(row["description"]))
            if score > best_score:
                best_row = row
                best_score = score
    if best_row and best_score >= 0.85:
        if not internal_part:
            line["internal_part_number"] = best_row["internal_part_number"]
        if not internal_revision and best_row["internal_part_revision"]:
            line["internal_part_revision"] = best_row["internal_part_revision"]
        return {
            "status": "matched_fuzzy",
            "product_id": best_row["id"],
            "canonical_product_id": canonical_product_id_for_legacy_product(conn, best_row["id"]),
            "customer_product_alias_id": alias["id"] if alias else None,
            "score": round(best_score, 3),
            "reason": f"Fuzzy description match to {best_row['internal_part_number']}.",
        }
    if best_row and best_score >= 0.55:
        return {
            "status": "needs_review",
            "product_id": best_row["id"],
            "canonical_product_id": canonical_product_id_for_legacy_product(conn, best_row["id"]),
            "customer_product_alias_id": alias["id"] if alias else None,
            "score": round(best_score, 3),
            "reason": f"Weak description match to {best_row['internal_part_number']}; review before using.",
        }
    return {
        "status": "unmatched",
        "product_id": None,
        "canonical_product_id": alias["product_id"] if alias else None,
        "customer_product_alias_id": alias["id"] if alias else None,
        "score": round(best_score, 3),
        "reason": "No product master match.",
    }


def product_tokens(value: Any) -> set[str]:
    tokens = re.findall(r"[a-z0-9]+", str(value or "").lower())
    return {token for token in tokens if len(token) > 1 and token not in COMMON_PRODUCT_WORDS}


def token_similarity(left: set[str], right: set[str]) -> float:
    if not left or not right:
        return 0.0
    overlap = len(left & right)
    union = len(left | right)
    score = overlap / union if union else 0.0
    numeric_left = {token for token in left if any(char.isdigit() for char in token)}
    numeric_right = {token for token in right if any(char.isdigit() for char in token)}
    if numeric_left and numeric_left & numeric_right:
        score += 0.15
    return min(score, 1.0)


def generate_review_tasks_for_po(conn: sqlite3.Connection, po_id: int) -> None:
    po = conn.execute("SELECT * FROM purchase_orders WHERE id = ?", (po_id,)).fetchone()
    if not po:
        return
    conn.execute(
        """
        UPDATE review_tasks
        SET status = 'resolved', resolved_at = CURRENT_TIMESTAMP, updated_at = CURRENT_TIMESTAMP
        WHERE purchase_order_id = ? AND status = 'open' AND created_by_system = 1
        """,
        (po_id,),
    )
    confidence = parse_json_dict(po["field_confidence_json"] if "field_confidence_json" in po.keys() else None)
    for field, label in [
        ("customer_company_name", "Customer company"),
        ("po_number", "PO number"),
        ("date_received", "Date received"),
    ]:
        if not po[field]:
            upsert_review_task(conn, po_id, None, "po_header", po_id, "missing_required_field", f"{label} is missing.", "critical", field, po[field], None, None)
    for field, value in confidence.items():
        try:
            score = float(value)
        except (TypeError, ValueError):
            continue
        if score < 0.7:
            upsert_review_task(conn, po_id, None, "po_header", po_id, "low_confidence", f"{field.replace('_', ' ').title()} has low extraction confidence.", "warning", field, po[field] if field in po.keys() else None, po[field] if field in po.keys() else None, score)
    for review in conn.execute("SELECT * FROM po_master_data_reviews WHERE purchase_order_id = ? AND status = 'open'", (po_id,)).fetchall():
        reason = {
            "customer": "unmatched_customer",
            "bill_to_address": "unmatched_bill_to",
            "ship_to_address": "unmatched_ship_to",
            "contact": "unmatched_contact",
        }.get(review["review_type"], "processing_error")
        upsert_review_task(conn, po_id, None, "master_data", review["id"], reason, review["message"], "warning", review["review_type"], None, None, None, review["suggested_value_json"])
    for line in conn.execute("SELECT * FROM purchase_order_lines WHERE purchase_order_id = ?", (po_id,)).fetchall():
        line_conf = parse_json_dict(line["field_confidence_json"] if "field_confidence_json" in line.keys() else None)
        if not line["customer_part_number"] and not line["description"]:
            upsert_review_task(conn, po_id, line["id"], "po_line", line["id"], "missing_required_field", "Line is missing both customer part number and description.", "critical", "customer_part_number", "", "", None)
        for field in ("quantity", "unit_price"):
            if line[field] in (None, ""):
                upsert_review_task(conn, po_id, line["id"], "po_line", line["id"], "missing_required_field", f"Line {line['line_number'] or line['id']} is missing {field.replace('_', ' ')}.", "warning", field, line[field], line[field], None)
        for field, value in line_conf.items():
            try:
                score = float(value)
            except (TypeError, ValueError):
                continue
            if score < 0.7:
                upsert_review_task(conn, po_id, line["id"], "po_line", line["id"], "low_confidence", f"Line {line['line_number'] or line['id']} {field.replace('_', ' ')} has low confidence.", "warning", field, line[field] if field in line.keys() else None, line[field] if field in line.keys() else None, score)
        if line["product_match_status"] in {"unmatched", "needs_review"}:
            reason = "weak_product_match" if line["product_match_status"] == "needs_review" else "unmatched_product"
            severity = "warning" if reason == "weak_product_match" else "critical"
            upsert_review_task(conn, po_id, line["id"], "product_match", line["id"], reason, line["product_match_reason"] or "Line item was not matched to product master.", severity, "internal_part_number", line["internal_part_number"], None, line["product_match_score"])
    if "differed from calculated line total" in (po["extraction_notes"] or "").lower():
        upsert_review_task(conn, po_id, None, "po_header", po_id, "total_mismatch", "Header total differs from calculated line total.", "warning", "total_value", po["total_value"], po["total_value"], None)
    for candidate in conn.execute("SELECT * FROM po_duplicate_candidates WHERE purchase_order_id = ? AND status = 'open'", (po_id,)).fetchall():
        reason_code = "duplicate_conflict" if candidate["match_type"] == "possible_duplicate" else "duplicate_po"
        severity = "critical" if candidate["match_type"] in {"same_po_number_revision", "same_attachment_hash", "possible_duplicate"} else "warning"
        upsert_review_task(
            conn,
            po_id,
            None,
            "duplicate",
            candidate["id"],
            reason_code,
            candidate["reason"] or "Possible duplicate purchase order found.",
            severity,
            "po_number",
            po["po_number"],
            po["po_number"],
            candidate["match_score"],
            json.dumps({"duplicate_candidate_id": candidate["id"], "candidate_purchase_order_id": candidate["candidate_purchase_order_id"], "match_type": candidate["match_type"]}),
        )
    conn.commit()


def upsert_review_task(
    conn: sqlite3.Connection,
    po_id: int,
    line_id: int | None,
    entity_type: str,
    entity_id: int | None,
    reason_code: str,
    message: str,
    severity: str,
    field_name: str | None,
    current_value: Any,
    extracted_value: Any,
    confidence: float | None,
    suggested_value_json: str | None = None,
    source_reference_json: str | None = None,
) -> None:
    ignored = conn.execute(
        """
        SELECT id FROM review_tasks
        WHERE purchase_order_id = ? AND COALESCE(purchase_order_line_id, 0) = COALESCE(?, 0)
          AND entity_type = ? AND reason_code = ? AND COALESCE(field_name, '') = COALESCE(?, '')
          AND status = 'ignored'
        """,
        (po_id, line_id, entity_type, reason_code, field_name),
    ).fetchone()
    if ignored:
        return
    existing = conn.execute(
        """
        SELECT id FROM review_tasks
        WHERE purchase_order_id = ? AND COALESCE(purchase_order_line_id, 0) = COALESCE(?, 0)
          AND entity_type = ? AND reason_code = ? AND COALESCE(field_name, '') = COALESCE(?, '')
          AND status = 'open'
        """,
        (po_id, line_id, entity_type, reason_code, field_name),
    ).fetchone()
    args = (
        message,
        severity,
        "" if current_value is None else str(current_value),
        "" if extracted_value is None else str(extracted_value),
        suggested_value_json,
        confidence,
        priority_for_severity(severity),
        source_reference_json,
    )
    if existing:
        conn.execute(
            """
            UPDATE review_tasks
            SET message = ?, severity = ?, current_value = ?, extracted_value = ?,
                suggested_value_json = ?, confidence = ?, priority = ?, source_reference_json = ?,
                last_seen_at = CURRENT_TIMESTAMP, updated_at = CURRENT_TIMESTAMP
            WHERE id = ?
            """,
            (*args, existing["id"]),
        )
        return
    conn.execute(
        """
        INSERT INTO review_tasks (
            purchase_order_id, purchase_order_line_id, entity_type, entity_id, reason_code, message,
            severity, status, field_name, current_value, extracted_value, suggested_value_json,
            confidence, created_by_system, priority, source_reference_json, last_seen_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, 'open', ?, ?, ?, ?, ?, 1, ?, ?, CURRENT_TIMESTAMP)
        """,
        (po_id, line_id, entity_type, entity_id, reason_code, message, severity, field_name, args[2], args[3], suggested_value_json, confidence, args[6], args[7]),
    )


def priority_for_severity(severity: str) -> int:
    return {"critical": 1, "warning": 2, "info": 3}.get(severity, 2)


def parse_json_dict(value: str | None) -> dict:
    try:
        parsed = json.loads(value or "{}")
        return parsed if isinstance(parsed, dict) else {}
    except json.JSONDecodeError:
        return {}
